import os
import time
import streamlit as st
import pandas as pd
from dotenv import load_dotenv
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# For PDFs, Word, Images
import docx
import pytesseract
from PIL import Image
import fitz  # PyMuPDF

# In Local
# load_dotenv()
# os.environ['GROQ_API_KEY'] = os.getenv("GROQ_API_KEY")
# os.environ['HF_TOKEN'] = os.getenv("HF_TOKEN")
# for streamlit
os.environ['GROQ_API_KEY'] = st.secrets["GROQ_API_KEY"]
os.environ['HF_TOKEN'] = st.secrets["HF_TOKEN"]


# Config
EXCEL_FILE = "./data/tickets.xlsx"
VECTOR_STORE_DIR = "vector_store"
LAST_MOD_TIME_FILE = "last_mod_time.txt"

# HuggingFace embeddings
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"}
)

# LLM
llm = ChatGroq(
    groq_api_key=os.environ["GROQ_API_KEY"],
    model_name="Llama3-8b-8192"
)

# Prompt template
prompt = ChatPromptTemplate.from_template("""
You are an IT support assistant. Use the ticket data below to answer the user's query.

<context>
{context}
</context>

Question: {input}

If there's no helpful information in the context, reply with:
"Not enough historical data to answer accurately."
""")

# Utility functions
def get_last_mod_time():
    return os.path.getmtime(EXCEL_FILE)

def read_last_saved_mod_time():
    if os.path.exists(LAST_MOD_TIME_FILE):
        with open(LAST_MOD_TIME_FILE, "r") as f:
            return float(f.read().strip())
    return None

def save_last_mod_time(mod_time):
    with open(LAST_MOD_TIME_FILE, "w") as f:
        f.write(str(mod_time))

def build_vector_store():
    df = pd.read_excel(EXCEL_FILE)

    if "Description" not in df.columns:
        st.error("❌ 'Description' column not found in Excel file.")
        st.stop()

    df.dropna(subset=["Description"], inplace=True)

    documents = []
    for _, row in df.iterrows():
        content = f"""
        Ticket ID: {row.get('Ticket ID', '')}
        Description: {row.get('Description', '')}
        Category: {row.get('Category', '')}
        Subcategory: {row.get('Subcategory', '')}
        Resolution: {row.get('Resolution', '')}
        """
        documents.append(Document(page_content=content.strip()))

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = text_splitter.split_documents(documents)

    vector_store = FAISS.from_documents(split_docs, embeddings)
    vector_store.save_local(VECTOR_STORE_DIR)
    st.success("✅ Vector DB rebuilt.")

def load_vector_store():
    return FAISS.load_local(VECTOR_STORE_DIR, embeddings, allow_dangerous_deserialization=True)

def extract_text_from_file(file):
    ext = file.name.lower()
    if ext.endswith(".txt"):
        return file.read().decode("utf-8")
    elif ext.endswith(".pdf"):
        pdf = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in pdf])
    elif ext.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif ext.endswith((".png", ".jpg", ".jpeg")):
        img = Image.open(file)
        return pytesseract.image_to_string(img)
    else:
        st.error("Unsupported file format.")
        return ""

# Streamlit UI
st.set_page_config(page_title="Ticket Gen AI Assistant", layout="wide")
st.title("🎫 Gen AI Ticket Assistant (Auto Refresh + Multi-Input)")

# Auto-rebuild vector store if Excel changed
last_mod_time = get_last_mod_time()
saved_mod_time = read_last_saved_mod_time()

if saved_mod_time != last_mod_time or not os.path.exists(VECTOR_STORE_DIR):
    st.warning("Excel file changed — rebuilding vector DB...")
    build_vector_store()
    save_last_mod_time(last_mod_time)
    st.stop()

# Load existing vector store
vector_store = load_vector_store()

# Query input options
query_type = st.radio("Choose input type:", ["Text", "File (PDF/DOCX/Image)"])

if query_type == "Text":
    user_query = st.text_input("📝 Describe a new ticket or ask a question:")
else:
    uploaded_query_file = st.file_uploader("📂 Upload a query file", type=["txt", "pdf", "docx", "png", "jpg", "jpeg"])
    user_query = extract_text_from_file(uploaded_query_file) if uploaded_query_file else ""

# Run query
if user_query:
    document_chain = create_stuff_documents_chain(llm, prompt)
    retriever = vector_store.as_retriever()
    retrieval_chain = create_retrieval_chain(retriever, document_chain)

    with st.spinner("Thinking..."):
        response = retrieval_chain.invoke({"input": user_query})
        st.write("### ✅ Answer")
        st.write(response['answer'])

        with st.expander("🔎 Similar Ticket Context"):
            for doc in response['context']:
                st.markdown(doc.page_content)
                st.write("---")
